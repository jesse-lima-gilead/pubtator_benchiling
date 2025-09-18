from pathlib import Path
from typing import Union, Dict, Any

from docx import Document


class DocxMetadataExtractor:
    def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        doc = Document(file_path)
        props = doc.core_properties
        meta: Dict[str, Any] = {
            "title": props.title,
            "author": props.author,
            "created": props.created,
            "modified": props.modified,
            "keywords": props.keywords,
            "subject": props.subject,
            "category": props.category,
            "revision": props.revision,
            "abstract": doc.paragraphs[0].text if doc.paragraphs else None,
            "sections": [
                p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
            ],
        }
        # Extract headings as section titles
        return meta
